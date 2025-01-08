import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from jobspy import scrape_jobs
import csv
from datetime import datetime
import pandas as pd
from langchain_ollama.llms import OllamaLLM
from langchain_core.prompts import ChatPromptTemplate
import PyPDF2
from docx import Document  # Updated import
import openpyxl
from openpyxl.styles import PatternFill

class JobSearchApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Job Search Tool")
        self.root.geometry("600x800")
        
        # Initialize resume content
        self.resume_content = None
        
        # Create main frame
        main_frame = ttk.Frame(root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Job Sites
        ttk.Label(main_frame, text="Job Sites:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.sites = ["indeed", "linkedin", "zip_recruiter", "glassdoor", "google"]
        self.site_vars = {}
        site_frame = ttk.Frame(main_frame)
        site_frame.grid(row=1, column=0, columnspan=2, sticky=tk.W, pady=5)
        
        for i, site in enumerate(self.sites):
            self.site_vars[site] = tk.BooleanVar(value=True)
            ttk.Checkbutton(site_frame, text=site, variable=self.site_vars[site]).grid(row=0, column=i, padx=5)
        
        # Search Term
        ttk.Label(main_frame, text="Search Term:").grid(row=2, column=0, sticky=tk.W, pady=5)
        self.search_term = ttk.Entry(main_frame, width=40)
        self.search_term.grid(row=2, column=1, sticky=tk.W, pady=5)
        
        # Location
        ttk.Label(main_frame, text="Location:").grid(row=3, column=0, sticky=tk.W, pady=5)
        self.location = ttk.Entry(main_frame, width=40)
        self.location.grid(row=3, column=1, sticky=tk.W, pady=5)
        
        # Job Type
        ttk.Label(main_frame, text="Job Type:").grid(row=4, column=0, sticky=tk.W, pady=5)
        self.job_type = ttk.Combobox(main_frame, values=["", "fulltime", "parttime", "internship", "contract"])
        self.job_type.grid(row=4, column=1, sticky=tk.W, pady=5)
        
        # Results Wanted
        ttk.Label(main_frame, text="Results per site:").grid(row=5, column=0, sticky=tk.W, pady=5)
        self.results_wanted = ttk.Entry(main_frame, width=10)
        self.results_wanted.insert(0, "20")
        self.results_wanted.grid(row=5, column=1, sticky=tk.W, pady=5)
        
        # Remote Option
        self.is_remote = tk.BooleanVar()
        ttk.Checkbutton(main_frame, text="Remote Only", variable=self.is_remote).grid(row=6, column=0, columnspan=2, sticky=tk.W, pady=5)
        
        # Country Selection
        ttk.Label(main_frame, text="Country (Indeed/Glassdoor):").grid(row=7, column=0, sticky=tk.W, pady=5)
        self.country = ttk.Combobox(main_frame, values=["USA", "UK", "Canada", "Australia", "Germany", "France"])
        self.country.set("USA")
        self.country.grid(row=7, column=1, sticky=tk.W, pady=5)
        
        # Resume Upload Section
        ttk.Label(main_frame, text="Resume:").grid(row=8, column=0, sticky=tk.W, pady=5)
        self.resume_label = ttk.Label(main_frame, text="No file selected")
        self.resume_label.grid(row=8, column=1, sticky=tk.W, pady=5)
        ttk.Button(main_frame, text="Upload Resume", command=self.upload_resume).grid(row=8, column=2, pady=5)
        
        # Progress Text
        self.progress_text = tk.Text(main_frame, height=10, width=60)
        self.progress_text.grid(row=9, column=0, columnspan=3, pady=10)
        
        # Search Button
        ttk.Button(main_frame, text="Search Jobs", command=self.search_jobs).grid(row=10, column=0, columnspan=3, pady=10)

    def log_progress(self, message):
        self.progress_text.insert(tk.END, f"{message}\n")
        self.progress_text.see(tk.END)
        self.root.update()

    def upload_resume(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("PDF files", "*.pdf"), ("Word files", "*.docx")]
        )
        if file_path:
            self.resume_content = self.parse_resume(file_path)
            self.resume_label.config(text=file_path.split("/")[-1])
            self.log_progress("Resume uploaded successfully")

    def parse_resume(self, file_path):
        """Parse resume content from PDF or DOCX file"""
        if file_path.endswith('.pdf'):
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
                return text
        elif file_path.endswith('.docx'):
            # Updated DOCX parsing using python-docx
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            # Also extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            return text
        return None

    def analyze_job_fit(self, job_description):
        """Use Ollama to analyze job fit based on resume"""
        if not self.resume_content:
            return "Not Apply (No Resume)"
        
        try:
            llm = OllamaLLM(model="llama3.2:latest")
            template = """
            Based on the following resume and job description, categorize this job as either:
            1. "Not Apply" (clear mismatches)
            2. "Should Apply" (good fit with some gaps)
            3. "Must Apply" (excellent fit)
            
            Resume:
            {resume}
            
            Job Description:
            {job_description}
            
            Respond with only one of the three categories listed above.
            """
            
            prompt = ChatPromptTemplate.from_template(template)
            chain = prompt | llm
            
            result = chain.invoke({
                "resume": self.resume_content,
                "job_description": job_description
            })
            
            return result.strip()
        except Exception as e:
            self.log_progress(f"Error in job analysis: {str(e)}")
            return "Analysis Error"

    def search_jobs(self):
        try:
            # Validate inputs and create parameters dict
            params = {
                'site': [site for site, var in self.site_vars.items() if var.get()],
                'search_term': self.search_term.get(),
                'location': self.location.get(),
                'results_wanted': int(self.results_wanted.get()),
                'country': self.country.get(),
            }
            
            if self.job_type.get():
                params['job_type'] = self.job_type.get()
            if self.is_remote.get():
                params['remote'] = True
                
            self.log_progress("Starting job search...")
            jobs = scrape_jobs(**params)
            
            if len(jobs) == 0:
                self.log_progress("No jobs found matching your criteria.")
                return
            
            # Create DataFrame with additional category column
            jobs_filtered = pd.DataFrame(columns=[
                'site', 'title', 'company', 'location', 'date_posted',
                'salary_min', 'salary_max', 'salary_interval', 'job_url',
                'description', 'category'
            ])
            
            # Analyze each job
            total_jobs = len(jobs)
            for index, job in jobs.iterrows():
                self.log_progress(f"Analyzing job {index + 1}/{total_jobs}")
                category = self.analyze_job_fit(job['description'])
                job['category'] = category
                jobs_filtered = pd.concat([jobs_filtered, pd.DataFrame([job])], ignore_index=True)
            
            # Generate filename with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"job_search_results_{timestamp}.xlsx"
            
            # Export to Excel with colors
            writer = pd.ExcelWriter(filename, engine='openpyxl')
            jobs_filtered.to_excel(writer, index=False, sheet_name='Jobs')
            
            # Apply conditional formatting
            workbook = writer.book
            worksheet = writer.sheets['Jobs']
            
            # Define colors for categories
            category_colors = {
                'Not Apply': 'FFB6C1',  # Light red
                'Should Apply': 'FFFACD',  # Light yellow
                'Must Apply': '90EE90'  # Light green
            }
            
            # Apply colors based on category
            category_col = jobs_filtered.columns.get_loc('category') + 1
            for row in range(2, len(jobs_filtered) + 2):
                category = worksheet.cell(row=row, column=category_col).value
                if category in category_colors:
                    fill = PatternFill(start_color=category_colors[category],
                                     end_color=category_colors[category],
                                     fill_type='solid')
                    for col in range(1, len(jobs_filtered.columns) + 1):
                        worksheet.cell(row=row, column=col).fill = fill
            
            writer.close()
            
            self.log_progress(f"Found {len(jobs_filtered)} jobs")
            self.log_progress(f"Results exported to: {filename}")
            
            # Show success message
            messagebox.showinfo("Success", 
                              f"Search completed!\nFound {len(jobs_filtered)} jobs\n"
                              f"Results saved to {filename}")
            
        except Exception as e:
            self.log_progress(f"Error: {str(e)}")
            messagebox.showerror("Error", f"An error occurred: {str(e)}")

def main():
    root = tk.Tk()
    
    # Center the window on the screen
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()
    window_width = 600
    window_height = 800  # Increased height for better layout
    x = (screen_width/2) - (window_width/2)
    y = (screen_height/2) - (window_height/2)
    root.geometry(f'{window_width}x{window_height}+{int(x)}+{int(y)}')
    
    app = JobSearchApp(root)
    root.mainloop()

if __name__ == "__main__":
    print("Starting application...")
    main()
    print("Application closed")