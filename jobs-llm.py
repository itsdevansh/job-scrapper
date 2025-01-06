import gradio as gr
import pandas as pd
import requests
from bs4 import BeautifulSoup
import PyPDF2
import docx
import json
from datetime import datetime, timedelta
import time
import re
from typing import Dict, List, Tuple, Optional
from openpyxl import Workbook
from openpyxl.styles import PatternFill
from langchain_core.prompts import ChatPromptTemplate
from langchain_ollama.llms import OllamaLLM
import numpy as np
from io import BytesIO
from requests.adapters import HTTPAdapter
from requests.packages.urllib3.util.retry import Retry
import logging
from fake_useragent import UserAgent
import html
import concurrent.futures
from functools import partial

class JobScraper():
    def __init__(self, debug_mode: bool = False):
        print("Initializing JobScraper...")
        
        # Set up logging
        logging.basicConfig(
            level=logging.DEBUG if debug_mode else logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        self.logger = logging.getLogger(__name__)
        print(f"Logging initialized with debug_mode: {debug_mode}")
        
        # Initialize session with retry strategy
        self.session = requests.Session()
        retries = Retry(
            total=5,
            backoff_factor=2,
            status_forcelist=[429, 500, 502, 503, 504],
            allowed_methods=["GET", "HEAD"]
        )
        self.session.mount('http://', HTTPAdapter(max_retries=retries))
        self.session.mount('https://', HTTPAdapter(max_retries=retries))
        print("Session initialized with retry strategy")
        
        self.ua = UserAgent()
        self.jobs_data = []
        self.debug_mode = debug_mode
        self.TIMEOUT = 15
        print("JobScraper initialization complete")

    def get_headers(self) -> dict:
        headers = {
            'User-Agent': self.ua.random,
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'DNT': '1',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1',
        }
        print(f"Generated headers with User-Agent: {headers['User-Agent'][:30]}...")
        return headers

    def safe_request(self, url: str, params: Optional[Dict] = None) -> Optional[requests.Response]:
        print(f"\nAttempting request to: {url}")
        print(f"With params: {params}")
        
        try:
            response = self.session.get(
                url,
                params=params,
                headers=self.get_headers(),
                timeout=self.TIMEOUT
            )
            print(f"Response status code: {response.status_code}")
            
            if response.status_code == 200:
                print("Request successful")
                return response
            
            print(f"Request failed with status code: {response.status_code}")
            return None
            
        except requests.RequestException as e:
            print(f"Request failed with error: {str(e)}")
            return None

    def scrape_linkedin(self, job_title: str, location: str = '', pages: int = 3) -> List[Dict]:
        print(f"\nStarting LinkedIn scraping for: {job_title} in {location}")
        print(f"Will scrape {pages} pages")
        
        base_url = 'https://www.linkedin.com/jobs/search'
        
        for page in range(pages):
            print(f"\nProcessing LinkedIn page {page + 1}/{pages}")
            
            search_params = {
                'keywords': job_title,
                'location': location,
                'start': page * 25,
                'sortBy': 'DD'
            }
            
            response = self.safe_request(base_url, search_params)
            if not response:
                print("Failed to get response, skipping page")
                continue
                
            try:
                print("Parsing page content...")
                soup = BeautifulSoup(response.text, 'html.parser')
                job_cards = soup.find_all('div', class_='job-search-card')
                
                print(f"Found {len(job_cards)} job cards")
                
                if not job_cards:
                    print(f"No job cards found on LinkedIn page {page + 1}")
                    continue
                
                for i, card in enumerate(job_cards, 1):
                    print(f"\nProcessing job card {i}/{len(job_cards)}")
                    try:
                        job_data = self._parse_linkedin_card(card)
                        if job_data:
                            print(f"Successfully parsed job: {job_data['title'][:30]}...")
                            self.jobs_data.append(job_data)
                        else:
                            print("Failed to parse job card")
                    except Exception as e:
                        print(f"Error parsing LinkedIn job card: {str(e)}")
                        continue
                
                delay = np.random.uniform(2, 4)
                print(f"Waiting {delay:.2f} seconds before next request...")
                time.sleep(delay)
                
            except Exception as e:
                print(f"Error processing LinkedIn page {page + 1}: {str(e)}")
                continue
        
        print(f"\nLinkedIn scraping complete. Total jobs collected: {len(self.jobs_data)}")
        return self.jobs_data

    def _parse_linkedin_card(self, card) -> Optional[Dict]:
        try:
            job_data = {
                'platform': 'LinkedIn',
                'title': self.clean_text(card.find('h3', class_='base-search-card__title').text if card.find('h3', class_='base-search-card__title') else 'N/A'),
                'company': self.clean_text(card.find('h4', class_='base-search-card__subtitle').text if card.find('h4', class_='base-search-card__subtitle') else 'N/A'),
                'location': self.clean_text(card.find('span', class_='job-search-card__location').text if card.find('span', class_='job-search-card__location') else 'N/A'),
                'date_posted': self.format_date(card.find('time')['datetime'] if card.find('time') else datetime.now().strftime('%Y-%m-%d')),
                'link': card.find('a', class_='base-card__full-link')['href'] if card.find('a', class_='base-card__full-link') else 'N/A',
                'salary': self.clean_text(card.find('span', class_='job-search-card__salary-info').text if card.find('span', class_='job-search-card__salary-info') else 'Not specified')
            }
            print(f"Successfully parsed job: {job_data['title']}")
            return job_data
            
        except Exception as e:
            print(f"Error parsing LinkedIn card: {str(e)}")
            return None

class JobMatcher:
    def __init__(self):
        print("\nInitializing JobMatcher...")
        
        try:
            print("Initializing Ollama LLM...")
            self.model = OllamaLLM(model="llama3.2:latest")
            print("LLM initialized successfully")
            
            # Initialize ResumeParser with the LLM client
            self.resume_parser = ResumeParser(llm_client=self.model)
            
            print("Setting up chat template...")
            self.template = ChatPromptTemplate.from_template("""
            Analyze the match between the resume and job description below.
            Focus on skills, experience, and qualifications match.
            
            Resume:
            {resume_text}
            
            Job Description:
            {job_description}
            
            Provide a JSON response with:
            1. Match percentage (0-100)
            2. Key matching skills
            3. Missing skills
            4. Overall assessment
            """)
            
            print("Testing Ollama connection...")
            test_response = self.chain.invoke({"resume_text": "test", "job_description": "test"})
            print("Ollama test successful")
            self.llm_available = True
            
        except Exception as e:
            print(f"Ollama initialization failed: {str(e)}")
            self.llm_available = False
            # Initialize ResumeParser without LLM client for fallback parsing
            self.resume_parser = ResumeParser()

import PyPDF2
import docx
from io import BytesIO
from markitdown import MarkItDown
from typing import Optional
import os
import tempfile

class ResumeParser:
    def __init__(self, llm_client=None):
        print("\nInitializing ResumeParser...")
        self.llm_client = llm_client
        if llm_client:
            print("Initializing MarkItDown with provided LLM client")
            self.md_converter = MarkItDown(llm_client=llm_client)
        else:
            print("No LLM client provided - will use traditional parsing methods")
            self.md_converter = None
            
    def parse_resume(self, file, file_extension: str) -> str:
        """
        Parse resume using either MarkItDown (if available) or traditional methods
        """
        print(f"\nParsing resume with extension: {file_extension}")
        
        if self.md_converter and file_extension.lower() in ('.pdf', '.docx', '.pptx', '.jpg', '.jpeg', '.png'):
            return self._parse_with_markitdown(file, file_extension)
        else:
            return self._parse_traditional(file, file_extension)
            
    def _parse_with_markitdown(self, file, file_extension: str) -> str:
        print("Using MarkItDown for parsing...")
        try:
            # Create a temporary file to save the binary content
            with tempfile.NamedTemporaryFile(suffix=file_extension, delete=False) as temp_file:
                if isinstance(file, bytes):
                    temp_file.write(file)
                else:
                    temp_file.write(file.read())
                temp_file_path = temp_file.name
            
            # Convert using MarkItDown
            result = self.md_converter.convert(temp_file_path)
            
            # Clean up temporary file
            os.unlink(temp_file_path)
            
            print("MarkItDown parsing successful")
            return result.text_content
            
        except Exception as e:
            print(f"Error in MarkItDown parsing: {str(e)}")
            print("Falling back to traditional parsing methods")
            return self._parse_traditional(file, file_extension)
            
    def _parse_traditional(self, file, file_extension: str) -> str:
        """
        Traditional parsing methods as fallback
        """
        print("Using traditional parsing methods...")
        try:
            if file_extension.lower() == '.pdf':
                return self._parse_pdf(file)
            elif file_extension.lower() == '.docx':
                return self._parse_docx(file)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            print(f"Error in traditional parsing: {str(e)}")
            raise
            
    def _parse_pdf(self, file) -> str:
        print("Parsing PDF using PyPDF2...")
        try:
            if isinstance(file, bytes):
                file_data = BytesIO(file)
            else:
                file_data = BytesIO(file.read())
                
            pdf_reader = PyPDF2.PdfReader(file_data)
            print(f"PDF has {len(pdf_reader.pages)} pages")
            
            text = ""
            for i, page in enumerate(pdf_reader.pages, 1):
                print(f"Processing page {i}/{len(pdf_reader.pages)}")
                text += page.extract_text()
            
            print(f"Successfully extracted {len(text)} characters")
            return text
            
        except Exception as e:
            print(f"Error parsing PDF: {str(e)}")
            raise
            
    def _parse_docx(self, file) -> str:
        print("Parsing DOCX using python-docx...")
        try:
            if isinstance(file, bytes):
                file_data = BytesIO(file)
            else:
                file_data = BytesIO(file.read())
                
            doc = docx.Document(file_data)
            print(f"DOCX has {len(doc.paragraphs)} paragraphs")
            
            text = ""
            for i, paragraph in enumerate(doc.paragraphs, 1):
                print(f"Processing paragraph {i}/{len(doc.paragraphs)}")
                text += paragraph.text + "\n"
            
            print(f"Successfully extracted {len(text)} characters")
            return text
            
        except Exception as e:
            print(f"Error parsing DOCX: {str(e)}")
            raise

def create_ui():
    print("\nInitializing Gradio UI...")
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s'
    )
    logger = logging.getLogger(__name__)

    def process_search(
        resume_file: gr.File,
        job_title: str,
        location: str,
        pages: int
    ) -> Tuple[str, Optional[pd.DataFrame]]:
        print("\nProcessing search request...")
        print(f"Job Title: {job_title}")
        print(f"Location: {location}")
        print(f"Pages to search: {pages}")
        
        try:
            print("Initializing JobMatcherUI...")
            matcher = JobMatcherUI()
            
            print("Executing search and matching...")
            status, results_df = matcher.search_and_match(
                resume_file,
                job_title,
                location,
                pages
            )
            
            if results_df is not None:
                print("Search successful, saving results...")
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f'job_matches_{timestamp}.xlsx'
                results_df.to_excel(filename, index=False)
                print(f"Results saved to {filename}")
                return f"Results saved to {filename}", results_df
            
            print(f"Search completed with status: {status}")
            return status, None
            
        except Exception as e:
            print(f"Error processing search: {str(e)}")
            return f"Error: {str(e)}", None

    print("Creating Gradio interface...")
    with gr.Blocks(title="Job Matcher") as app:
        gr.Markdown("# Job Matcher")
        gr.Markdown("Upload your resume and search for matching jobs")
        
        with gr.Row():
            with gr.Column():
                resume_file = gr.File(
                    label="Upload Resume (PDF or DOCX)",
                    file_types=[".pdf", ".docx"]
                )
                job_title = gr.Textbox(
                    label="Job Title",
                    placeholder="e.g. Software Engineer"
                )
                location = gr.Textbox(
                    label="Location",
                    placeholder="e.g. San Francisco, CA"
                )
                pages = gr.Slider(
                    minimum=1,
                    maximum=10,
                    value=3,
                    step=1,
                    label="Number of Pages to Search"
                )
                search_button = gr.Button("Search and Match Jobs")
        
        with gr.Row():
            status_output = gr.Textbox(label="Status")
            results_table = gr.DataFrame(label="Job Matches")
        
        search_button.click(
            fn=process_search,
            inputs=[resume_file, job_title, location, pages],
            outputs=[status_output, results_table]
        )
    
    print("Gradio interface created successfully")
    return app

if __name__ == "__main__":
    print("\nStarting Job Matcher application...")
    app = create_ui()
    print("Launching Gradio server...")
    app.launch(
        share=True,
        server_name="0.0.0.0",
        server_port=7860,
        show_error=True
    )